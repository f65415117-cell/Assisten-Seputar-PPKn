import streamlit as st
import google.generativeai as genai
from docx import Document
from io import BytesIO

# 1. KONFIGURASI HALAMAN
st.set_page_config(page_title="Seputar PPKn AI", layout="centered")

# 2. FUNGSI DOWNLOAD WORD (Mendukung Teks Hasil AI)
def to_word(text):
    doc = Document()
    doc.add_heading('Hasil Soal & Kisi-kisi - Seputar PPKn AI', 0)
    doc.add_paragraph(text)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# 3. CSS MINIMALIS & SIMETRIS (ALA RUMAH PENDIDIKAN)
st.markdown("""
    <style>
    .stApp { background-color: #ffffff; }
    .header-minimalis { 
        display: flex; 
        align-items: center; 
        padding: 20px 0px; 
        border-bottom: 1px solid #eaeaea; 
        margin-bottom: 30px; 
    }
    .logo-img { 
        border-radius: 50%; 
        margin-right: 15px; 
        border: 2px solid #007bff; 
        padding: 2px;
    }
    .title-text { color: #333; font-size: 1.6rem; font-weight: 800; margin: 0; line-height: 1.2; }
    .subtitle-text { color: #666; font-size: 0.95rem; margin: 0; }
    
    /* Tombol Biru Gradasi */
    .stButton>button { 
        width: 100%; 
        background: linear-gradient(135deg, #007bff 0%, #0056b3 100%); 
        color: white; 
        border-radius: 10px; 
        font-weight: bold; 
        height: 3.5em; 
        border: none; 
        box-shadow: 0 4px 6px rgba(0,0,0,0.1);
        transition: 0.3s;
    }
    .stButton>button:hover { transform: translateY(-2px); box-shadow: 0 6px 12px rgba(0,0,0,0.15); }
    
    /* Merapikan Jarak Form */
    [data-testid="stVerticalBlock"] > div { padding-bottom: 0px; }
    </style>
    """, unsafe_allow_html=True)

# --- KONEKSI API ---
if "GOOGLE_API_KEY" in st.secrets:
    genai.configure(api_key=st.secrets["GOOGLE_API_KEY"])
else:
    st.error("API Key belum terpasang di Secrets!")

# 4. HEADER DENGAN LOGO SPN PPKn
logo_url = "https://yt3.googleusercontent.com/ytc/AIdro_k9jAOBysirU8tWHJ6xT4OQs6OvIBkC7JIjXf5uiUPKuA=s900-c-k-c0x00ffffff-no-rj"

st.markdown(f"""
    <div class="header-minimalis">
        <img src="{logo_url}" class="logo-img" width="65">
        <div>
            <h1 class="title-text">Seputar PPKn AI</h1>
            <p class="subtitle-text">Asisten Pembelajaran Digital PPKn</p>
        </div>
    </div>
    """, unsafe_allow_html=True)

# 5. FORM INPUT (VERSI SIMETRIS & RAPAT)
with st.container():
    col1, col2 = st.columns(2)
    with col1:
        kelas = st.selectbox("1. Pilih Kelas", ["Pilih...", "Kelas VII", "Kelas VIII", "Kelas IX", "Kelas X", "Kelas XI", "Kelas XII"])
    with col2:
        jenis = st.selectbox("2. Jenis Soal", ["Pilih...", "Pilihan Ganda", "Esai HOTS", "Menjodohkan"])
    
    col3, col4 = st.columns(2)
    with col3:
        level = st.selectbox("3. Level Kognitif", ["Pilih...", "C1-C2 (Pemahaman)", "C3-C4 (Aplikasi/Analisis)", "C5-C6 (Evaluasi/Kreasi)"])
    with col4:
        jumlah = st.number_input("4. Jumlah Soal", 1, 50, 5)
    
    topik = st.text_area("5. Topik, Bab, atau Kisi-kisi Soal:", 
                         placeholder="Contoh: Kedaulatan NKRI... (Bisa paste kisi-kisi soal di sini agar hasil lebih akurat)",
                         height=120)

    # Tombol Generate
    st.markdown("<div style='margin-top: 10px;'></div>", unsafe_allow_html=True)
    if st.button("🚀 GENERATE SOAL & KISI-KISI"):
        if kelas == "Pilih..." or jenis == "Pilih..." or not topik:
            st.warning("Lengkapi data Kelas, Jenis, dan Topik dulu ya, Bro!")
        else:
            try:
                # MODEL VERSI PILIHANMU
                model = genai.GenerativeModel("gemini-2.5-flash-lite")
                
                prompt = f"""
                Bertindaklah sebagai Pakar Kurikulum PPKn Indonesia.
                Tugas: Buatlah Kisi-kisi Soal dan daftar soal {jenis} untuk {kelas} dengan materi {topik}.
                Level kognitif yang diinginkan: {level}.

                STRUKTUR OUTPUT WAJIB:
                1. TABEL KISI-KISI (No, Lingkup Materi, Indikator Soal, Level, No Soal).
                2. DAFTAR SOAL (Opsi A, B, C, D HARUS ditulis berderet ke bawah).
                3. KUNCI JAWABAN DAN PEMBAHASAN.

                Gunakan standar Kurikulum Merdeka terbaru.
                """
                
                with st.spinner("Gemini 2.5 Flash Lite sedang menyusun administrasi lengkap..."):
                    response = model.generate_content(prompt)
                    hasil_akhir = response.text
                    
                    st.markdown("### 📝 Hasil Analisis:")
                    st.write(hasil_akhir)
                    
                    st.download_button(
                        label="📥 Download File Word (Soal + Kisi-kisi)",
                        data=to_word(hasil_akhir),
                        file_name=f"Administrasi_Soal_{topik}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
            except Exception as e:
                st.error(f"Terjadi kendala: {e}")
